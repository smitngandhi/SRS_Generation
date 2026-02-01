"""
SRS Document Generator Utility with Table of Contents

This module provides functionality to generate professionally formatted 
Software Requirements Specification (SRS) documents from JSON data with
automatic Table of Contents and page numbering.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, Any, List, Optional


class SRSDocumentGenerator:
    """Generate SRS documents from JSON data with proper formatting and TOC."""
    
    def __init__(self, project_name: str, authors: List[str] = None, organization: str = "Organization Name"):
        """
        Initialize the SRS document generator.
        
        Args:
            project_name: Name of the project for headers
            authors: List of document author names (default: ["Author Name"])
            organization: Organization name (default: "Organization Name")
        """
        self.project_name = project_name
        self.authors = authors if authors else ["Author Name"]
        self.organization = organization
        self.doc = Document()
        self._setup_document()
        self._setup_styles()
        
    def _setup_document(self):
        """Configure document-level settings (margins, page size)."""
        sections = self.doc.sections
        for section in sections:
            # Set margins: 1.25" left, 1" top/right/bottom
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            
            # Set page size to US Letter
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            
    def _setup_styles(self):
        """Set up custom styles for the document."""
        styles = self.doc.styles
        
        # Modify Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(12)
        
        # Heading 1 style
        try:
            h1_style = styles['Heading 1']
        except KeyError:
            h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Arial'
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 0, 0)
        h1_style.paragraph_format.space_before = Pt(24)
        h1_style.paragraph_format.space_after = Pt(12)
        
        # Heading 2 style
        try:
            h2_style = styles['Heading 2']
        except KeyError:
            h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Arial'
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0, 0, 0)
        h2_style.paragraph_format.space_before = Pt(18)
        h2_style.paragraph_format.space_after = Pt(9)
        
        # Heading 3 style
        try:
            h3_style = styles['Heading 3']
        except KeyError:
            h3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Arial'
        h3_font.size = Pt(13)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(0, 0, 0)
        h3_style.paragraph_format.space_before = Pt(12)
        h3_style.paragraph_format.space_after = Pt(6)
        
        # TOC Heading style
        try:
            toc_heading_style = styles['TOC Heading']
        except KeyError:
            toc_heading_style = styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
        toc_heading_font = toc_heading_style.font
        toc_heading_font.name = 'Arial'
        toc_heading_font.size = Pt(16)
        toc_heading_font.bold = True
        toc_heading_font.color.rgb = RGBColor(0, 0, 0)
        toc_heading_style.paragraph_format.space_before = Pt(24)
        toc_heading_style.paragraph_format.space_after = Pt(12)
        toc_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def _add_header_footer(self):
        """Add header and footer to all content pages (not title/TOC)."""
        # Get the last section (content section)
        sections = self.doc.sections
        
        # Configure all sections
        for idx, section in enumerate(sections):
            if idx == 0:
                # Title page section - no header/footer page numbers
                continue
            elif idx == 1:
                # TOC page section - no header/footer page numbers (optional)
                continue
            else:
                # Content sections - add header and footer
                # Header
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.text = f"SRS for {self.project_name}"
                header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if header_para.runs:
                    header_para.runs[0].font.size = Pt(10)
                    header_para.runs[0].font.name = 'Arial'
                
                # Footer with page number
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Clear existing content
                footer_para.clear()
                
                # Add page number field
                run = footer_para.add_run()
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = "PAGE"
                
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
                run.font.size = Pt(10)
                run.font.name = 'Arial'
        
    def _add_title_page(self):
        """Add title page for the SRS document."""
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Software Requirements Specification")
        run.font.name = 'Arial'
        run.font.size = Pt(20)
        run.font.bold = True
        
        # Add spacing
        title.paragraph_format.space_after = Pt(12)
        
        # "for" text
        for_text = self.doc.add_paragraph()
        for_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = for_text.add_run("for")
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        
        for_text.paragraph_format.space_after = Pt(12)
        
        # Project name
        project = self.doc.add_paragraph()
        project.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = project.add_run(self.project_name)
        run.font.name = 'Arial'
        run.font.size = Pt(20)
        run.font.bold = True
        
        project.paragraph_format.space_after = Pt(36)
        
        # Prepared by - format authors list
        prepared_by = self.doc.add_paragraph()
        prepared_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
        authors_text = ", ".join(self.authors)
        run = prepared_by.add_run(f"Prepared by {authors_text}")
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        
        prepared_by.paragraph_format.space_after = Pt(12)
        
        # Organization
        organization = self.doc.add_paragraph()
        organization.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = organization.add_run(f"Organization: {self.organization}")
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        
        organization.paragraph_format.space_after = Pt(12)
        
        # Date created
        from datetime import datetime
        date_created = self.doc.add_paragraph()
        date_created.alignment = WD_ALIGN_PARAGRAPH.CENTER
        today_date = datetime.now().strftime("%m/%d/%Y")
        run = date_created.add_run(f"Date Created: {today_date}")
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        
        # Add section break (new page) after title page
        self.doc.add_section()
    
    def _add_table_of_contents(self):
        """Add Table of Contents after the title page."""
        # Add TOC heading
        toc_heading = self.doc.add_paragraph()
        toc_heading.style = 'TOC Heading'
        run = toc_heading.add_run("Table of Contents")
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_heading.paragraph_format.space_after = Pt(18)
        
        # Add TOC field with proper structure
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        
        # Create the field start
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        
        # Create the instruction text
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        # Create the field separator
        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        
        # Add the field elements to the run
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_separate)
        
        # Add placeholder text (will be replaced when field is updated)
        run._r.append(OxmlElement('w:t'))
        
        # Create the field end
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_end)
        
        # Set update fields on open flag in document settings
        self._set_update_fields_on_open()
        
        # Add instructional text
        instruction = self.doc.add_paragraph()
        instruction_run = instruction.add_run(
            "\n[Note: The Table of Contents will be automatically generated when you open this document in Microsoft Word. "
            "If it doesn't appear, right-click and select 'Update Field']"
        )
        instruction_run.font.italic = True
        instruction_run.font.size = Pt(10)
        instruction_run.font.color.rgb = RGBColor(128, 128, 128)
        instruction.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add section break (new page) after TOC and restart page numbering
        new_section = self.doc.add_section()
        
        # Set page numbering to start at 1 for content section
        new_section.start_type = 2  # New page
        new_section.page_number_start = 1
    
    def _set_update_fields_on_open(self):
        """Set the document to update fields when opened."""
        try:
            # Access document settings
            settings_element = self.doc.settings.element
            
            # Create updateFields element if it doesn't exist
            update_fields = settings_element.find(qn('w:updateFields'))
            if update_fields is None:
                update_fields = OxmlElement('w:updateFields')
                update_fields.set(qn('w:val'), 'true')
                settings_element.append(update_fields)
            else:
                update_fields.set(qn('w:val'), 'true')
        except Exception as e:
            # If settings don't exist, create them
            pass
        
    def add_introduction_section(self, intro_data: Dict[str, Any]):
        """
        Add Introduction section to the document.
        
        Args:
            intro_data: Dictionary containing introduction section data
        """
        # Section title
        self.doc.add_heading(intro_data.get('title', '1. Introduction'), level=1)
        
        # 1.1 Purpose
        purpose = intro_data.get('purpose', {})
        self.doc.add_heading(purpose.get('title', '1.1 Purpose'), level=2)
        self.doc.add_paragraph(purpose.get('description', ''))
        
        # 1.2 Intended Audience
        audience = intro_data.get('intended_audience', {})
        self.doc.add_heading(audience.get('title', '1.2 Intended Audience'), level=2)
        audience_groups = audience.get('audience_groups', [])
        if audience_groups:
            para = self.doc.add_paragraph("This document is intended for:")
            for group in audience_groups:
                self.doc.add_paragraph(group, style='List Bullet')
        
        # 1.3 Project Scope
        scope = intro_data.get('project_scope', {})
        self.doc.add_heading(scope.get('title', '1.3 Project Scope'), level=2)
        
        included = scope.get('included', [])
        if included:
            self.doc.add_paragraph("Included in scope:", style='Heading 3')
            for item in included:
                self.doc.add_paragraph(item, style='List Bullet')
        
        excluded = scope.get('excluded', [])
        if excluded:
            self.doc.add_paragraph("Excluded from scope:", style='Heading 3')
            for item in excluded:
                self.doc.add_paragraph(item, style='List Bullet')
        
        # 1.4 Document Conventions
        conventions = intro_data.get('document_conventions', {})
        self.doc.add_heading(conventions.get('title', '1.4 Document Conventions'), level=2)
        conv_list = conventions.get('conventions', [])
        for conv in conv_list:
            self.doc.add_paragraph(conv, style='List Bullet')
        
        # 1.5 References
        references = intro_data.get('references', {})
        self.doc.add_heading(references.get('title', '1.5 References'), level=2)
        ref_list = references.get('references', [])
        for ref in ref_list:
            ref_id = ref.get('id', '')
            ref_desc = ref.get('description', '')
            self.doc.add_paragraph(f"{ref_id}: {ref_desc}", style='List Bullet')
    
    def add_overall_description_section(self, desc_data: Dict[str, Any]):
        """
        Add Overall Description section to the document.
        
        Args:
            desc_data: Dictionary containing overall description section data
        """
        # Section title
        self.doc.add_heading(desc_data.get('title', '2. Overall Description'), level=1)
        
        # 2.1 Product Perspective
        perspective = desc_data.get('product_perspective', {})
        self.doc.add_heading(perspective.get('title', '2.1 Product Perspective'), level=2)
        self.doc.add_paragraph(perspective.get('description', ''))
        
        # 2.2 Product Features
        features = desc_data.get('product_features', {})
        self.doc.add_heading(features.get('title', '2.2 Product Features'), level=2)
        feature_list = features.get('features', [])
        for feature in feature_list:
            self.doc.add_paragraph(feature, style='List Bullet')
        
        # 2.3 User Classes and Characteristics
        user_classes = desc_data.get('user_classes_and_characteristics', {})
        self.doc.add_heading(user_classes.get('title', '2.3 User Classes and Characteristics'), level=2)
        classes = user_classes.get('user_classes', [])
        for user_class in classes:
            user_type = user_class.get('user_class', '')
            self.doc.add_paragraph(user_type, style='Heading 3')
            characteristics = user_class.get('characteristics', [])
            for char in characteristics:
                self.doc.add_paragraph(char, style='List Bullet')
        
        # 2.4 Operating Environment
        environment = desc_data.get('operating_environment', {})
        self.doc.add_heading(environment.get('title', '2.4 Operating Environment'), level=2)
        env_list = environment.get('environments', [])
        for env in env_list:
            self.doc.add_paragraph(env, style='List Bullet')
        
        # 2.5 Design and Implementation Constraints
        constraints = desc_data.get('design_and_implementation_constraints', {})
        self.doc.add_heading(constraints.get('title', '2.5 Design and Implementation Constraints'), level=2)
        constraint_list = constraints.get('constraints', [])
        for constraint in constraint_list:
            self.doc.add_paragraph(constraint, style='List Bullet')
        
        # 2.6 User Documentation
        documentation = desc_data.get('user_documentation', {})
        self.doc.add_heading(documentation.get('title', '2.6 User Documentation'), level=2)
        doc_list = documentation.get('documents', [])
        for doc in doc_list:
            self.doc.add_paragraph(doc, style='List Bullet')
        
        # 2.7 Assumptions and Dependencies
        assumptions = desc_data.get('assumptions_and_dependencies', {})
        self.doc.add_heading(assumptions.get('title', '2.7 Assumptions and Dependencies'), level=2)
        
        assumption_list = assumptions.get('assumptions', [])
        if assumption_list:
            self.doc.add_paragraph("Assumptions:", style='Heading 3')
            for assumption in assumption_list:
                self.doc.add_paragraph(assumption, style='List Bullet')
        
        dependency_list = assumptions.get('dependencies', [])
        if dependency_list:
            self.doc.add_paragraph("Dependencies:", style='Heading 3')
            for dependency in dependency_list:
                self.doc.add_paragraph(dependency, style='List Bullet')
    
    def add_system_features_section(self, features_data: Dict[str, Any]):
        """
        Add System Features section to the document.
        
        Args:
            features_data: Dictionary containing system features section data
        """
        # Section title
        self.doc.add_heading(f"3. {features_data.get('title', 'System Features')}", level=1)
        
        features = features_data.get('features', [])
        for idx, feature in enumerate(features, 1):
            feature_name = feature.get('feature_name', f'Feature {idx}')
            self.doc.add_heading(f"3.{idx} {feature_name}", level=2)
            
            # Description
            description = feature.get('description', '')
            if description:
                self.doc.add_paragraph(f"Description: {description}")
            
            # Stimulus/Response Sequences
            stimulus_response = feature.get('stimulus_response', [])
            if stimulus_response:
                self.doc.add_paragraph("Stimulus/Response Sequences:", style='Heading 3')
                for sr in stimulus_response:
                    stimulus = sr.get('stimulus', '')
                    response = sr.get('response', '')
                    para = self.doc.add_paragraph()
                    para.add_run("Stimulus: ").bold = True
                    para.add_run(stimulus)
                    para = self.doc.add_paragraph()
                    para.add_run("Response: ").bold = True
                    para.add_run(response)
            
            # Functional Requirements
            functional_reqs = feature.get('functional_requirements', [])
            if functional_reqs:
                self.doc.add_paragraph("Functional Requirements:", style='Heading 3')
                for req in functional_reqs:
                    req_desc = req.get('description', '')
                    self.doc.add_paragraph(req_desc, style='List Bullet')
    
    def add_external_interfaces_section(
        self, 
        interfaces_data: Dict[str, Any],
        image_paths: Dict[str, str]
    ):
        """
        Add External Interface Requirements section to the document.
        
        Args:
            interfaces_data: Dictionary containing external interfaces section data
            image_paths: Dictionary with paths to interface diagrams
        """
        # Section title
        self.doc.add_heading(interfaces_data.get('title', '4. External Interface Requirements'), level=1)
        
        # 4.1 User Interfaces
        user_interfaces = interfaces_data.get('user_interfaces', {})
        self.doc.add_heading(user_interfaces.get('title', '4.1 User Interfaces'), level=2)
        self.doc.add_paragraph(user_interfaces.get('description', ''))
        
        # Add user interface diagram
        if 'user_interfaces' in image_paths and image_paths['user_interfaces'] and Path(image_paths['user_interfaces']).exists():
            self.doc.add_paragraph("User Interface Architecture:", style='Heading 3')
            self.doc.add_picture(str(image_paths['user_interfaces']), width=Inches(6.0))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 4.2 Hardware Interfaces
        hardware_interfaces = interfaces_data.get('hardware_interfaces', {})
        self.doc.add_heading(hardware_interfaces.get('title', '4.2 Hardware Interfaces'), level=2)
        self.doc.add_paragraph(hardware_interfaces.get('description', ''))
        
        # Add hardware interface diagram
        if 'hardware_interfaces' in image_paths and image_paths['hardware_interfaces'] and Path(image_paths['hardware_interfaces']).exists():
            self.doc.add_paragraph("Hardware Interface Architecture:", style='Heading 3')
            self.doc.add_picture(str(image_paths['hardware_interfaces']), width=Inches(6.0))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 4.3 Software Interfaces
        software_interfaces = interfaces_data.get('software_interfaces', {})
        self.doc.add_heading(software_interfaces.get('title', '4.3 Software Interfaces'), level=2)
        self.doc.add_paragraph(software_interfaces.get('description', ''))
        
        # Add software interface diagram
        if 'software_interfaces' in image_paths and image_paths['software_interfaces'] and Path(image_paths['software_interfaces']).exists():
            self.doc.add_paragraph("Software Interface Architecture:", style='Heading 3')
            self.doc.add_picture(str(image_paths['software_interfaces']), width=Inches(6.0))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 4.4 Communication Interfaces
        communication_interfaces = interfaces_data.get('communication_interfaces', {})
        self.doc.add_heading(communication_interfaces.get('title', '4.4 Communication Interfaces'), level=2)
        self.doc.add_paragraph(communication_interfaces.get('description', ''))
        
        # Add communication interface diagram
        if 'communication_interfaces' in image_paths and image_paths['communication_interfaces'] and Path(image_paths['communication_interfaces']).exists():
            self.doc.add_paragraph("Communication Interface Architecture:", style='Heading 3')
            self.doc.add_picture(str(image_paths['communication_interfaces']), width=Inches(6.0))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_nfr_section(self, nfr_data: Dict[str, Any]):
        """
        Add Non-Functional Requirements section to the document.
        
        Args:
            nfr_data: Dictionary containing NFR section data
        """
        # Section title
        self.doc.add_heading(f"5. {nfr_data.get('title', 'Non-Functional Requirements')}", level=1)
        
        # Performance Requirements
        performance = nfr_data.get('performance_requirements', {})
        self.doc.add_heading(f"5.1 {performance.get('title', 'Performance Requirements')}", level=2)
        perf_reqs = performance.get('requirements', [])
        for req in perf_reqs:
            desc = req.get('description', '')
            rationale = req.get('rationale', '')
            para = self.doc.add_paragraph(desc, style='List Bullet')
            if rationale:
                para.add_run(f"\nRationale: {rationale}").italic = True
        
        # Safety Requirements
        safety = nfr_data.get('safety_requirements', {})
        self.doc.add_heading(f"5.2 {safety.get('title', 'Safety Requirements')}", level=2)
        safety_reqs = safety.get('requirements', [])
        for req in safety_reqs:
            desc = req.get('description', '')
            rationale = req.get('rationale', '')
            para = self.doc.add_paragraph(desc, style='List Bullet')
            if rationale:
                para.add_run(f"\nRationale: {rationale}").italic = True
        
        # Security Requirements
        security = nfr_data.get('security_requirements', {})
        self.doc.add_heading(f"5.3 {security.get('title', 'Security Requirements')}", level=2)
        security_reqs = security.get('requirements', [])
        for req in security_reqs:
            desc = req.get('description', '')
            rationale = req.get('rationale', '')
            para = self.doc.add_paragraph(desc, style='List Bullet')
            if rationale:
                para.add_run(f"\nRationale: {rationale}").italic = True
        
        # Quality Attributes
        quality = nfr_data.get('quality_attributes', {})
        self.doc.add_heading(f"5.4 {quality.get('title', 'Quality Attributes')}", level=2)
        quality_reqs = quality.get('requirements', [])
        for req in quality_reqs:
            desc = req.get('description', '')
            rationale = req.get('rationale', '')
            para = self.doc.add_paragraph(desc, style='List Bullet')
            if rationale:
                para.add_run(f"\nRationale: {rationale}").italic = True
    
    def add_glossary_section(self, glossary_data: Dict[str, Any]):
        """
        Add Glossary section to the document.
        
        Args:
            glossary_data: Dictionary containing glossary section data
        """
        # Section title
        self.doc.add_heading("6. Glossary", level=1)
        
        sections = glossary_data.get('sections', [])
        for idx, section in enumerate(sections, 1):
            section_title = section.get('title', '')
            self.doc.add_heading(f"6.{idx} {section_title}", level=2)
            
            terms = section.get('terms', [])
            for term_data in terms:
                term = term_data.get('term', '')
                definition = term_data.get('definition', '')
                para = self.doc.add_paragraph()
                para.add_run(f"{term}: ").bold = True
                para.add_run(definition)
    
    def add_assumptions_section(self, assumptions_data: Dict[str, Any]):
        """
        Add Assumptions section to the document.
        
        Args:
            assumptions_data: Dictionary containing assumptions section data
        """
        # Section title
        self.doc.add_heading(f"7. {assumptions_data.get('title', 'Assumptions')}", level=1)
        
        assumptions = assumptions_data.get('assumptions', [])
        for idx, assumption in enumerate(assumptions, 1):
            description = assumption.get('description', '')
            impact = assumption.get('impact', '')
            
            self.doc.add_paragraph(f"Assumption {idx}:", style='Heading 3')
            self.doc.add_paragraph(description)
            
            if impact:
                para = self.doc.add_paragraph()
                para.add_run("Impact: ").bold = True
                para.add_run(impact)
    
    def save(self, output_path: str):
        """
        Save the document to the specified path.
        
        Args:
            output_path: Path where the document should be saved
        """
        self.doc.save(output_path)


def generate_srs_document(
    project_name: str,
    introduction_section: Dict[str, Any],
    overall_description_section: Dict[str, Any],
    system_features_section: Dict[str, Any],
    external_interfaces_section: Dict[str, Any],
    nfr_section: Dict[str, Any],
    glossary_section: Dict[str, Any],
    assumptions_section: Dict[str, Any],
    image_paths: Dict[str, str],
    output_path: str,
    authors: List[str] = None,
    organization: str = "Organization Name"
) -> str:
    """
    Generate a complete SRS document from JSON data with Table of Contents.
    
    Args:
        project_name: Name of the project
        introduction_section: Introduction section data
        overall_description_section: Overall description section data
        system_features_section: System features section data
        external_interfaces_section: External interfaces section data
        nfr_section: Non-functional requirements section data
        glossary_section: Glossary section data
        assumptions_section: Assumptions section data
        image_paths: Dictionary with paths to interface diagrams
            Expected keys: 'user_interfaces', 'hardware_interfaces', 
                          'software_interfaces', 'communication_interfaces'
        output_path: Path where the document should be saved
        authors: List of document author names (default: ["Author Name"])
        organization: Organization name (default: "Organization Name")
    
    Returns:
        str: Path to the generated document
    
    Example:
        ```python
        image_paths = {
            'user_interfaces': './static/HireSmart_user_interfaces_diagram.png',
            'hardware_interfaces': './static/HireSmart_hardware_interfaces_diagram.png',
            'software_interfaces': './static/HireSmart_software_interfaces_diagram.png',
            'communication_interfaces': './static/HireSmart_communication_interfaces_diagram.png'
        }
        
        generate_srs_document(
            project_name="HireSmart",
            introduction_section=intro_data,
            overall_description_section=desc_data,
            system_features_section=features_data,
            external_interfaces_section=interfaces_data,
            nfr_section=nfr_data,
            glossary_section=glossary_data,
            assumptions_section=assumptions_data,
            image_paths=image_paths,
            output_path="./output/HireSmart_SRS.docx",
            authors=["John Doe", "Jane Smith"],
            organization="ABC Corporation"
        )
        ```
    """
    # Create generator instance
    generator = SRSDocumentGenerator(project_name, authors, organization)
    
    # Add title page
    generator._add_title_page()
    
    # Add Table of Contents
    generator._add_table_of_contents()
    
    # Add all sections
    generator.add_introduction_section(introduction_section)
    generator.add_overall_description_section(overall_description_section)
    generator.add_system_features_section(system_features_section)
    generator.add_external_interfaces_section(external_interfaces_section, image_paths)
    generator.add_nfr_section(nfr_section)
    generator.add_glossary_section(glossary_section)
    generator.add_assumptions_section(assumptions_section)
    
    # Add header and footer (must be after all content is added)
    generator._add_header_footer()
    
    # Save the document
    generator.save(output_path)
    
    return output_path